import json
import time
import base64
import re
import io
import zipfile
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional

import numpy as np
from fastapi import FastAPI, UploadFile, File, Form
from openai import OpenAI
from monitoring import setup_monitoring
import os
from datetime import datetime

# =========================
# LangSmith Setup
# =========================
from dotenv import load_dotenv
from langsmith import traceable

load_dotenv()
os.environ.setdefault("LANGSMITH_TRACING", "true")
os.environ.setdefault("LANGSMITH_PROJECT", "LLMProject")
os.environ.setdefault("LANGSMITH_ENDPOINT", "https://api.smith.langchain.com")


app = FastAPI(title="Scam Phishing Hybrid API")

client = OpenAI()

BASE_DIR = Path(__file__).resolve().parent
EMBEDDINGS_PATH = BASE_DIR / "data" / "phishing_embeddings.json"

logger = setup_monitoring(app, BASE_DIR)

# =========================
# Chat History Logging (FOR EVALUATION)
# =========================

CHAT_HISTORY_PATH = BASE_DIR / "logs" / "chat_history.jsonl"

def save_chat_history(user_message, bot_answer, mode="unknown"):
    try:
        os.makedirs(CHAT_HISTORY_PATH.parent, exist_ok=True)

        record = {
            "timestamp": datetime.now().isoformat(),
            "user_message": str(user_message),
            "bot_answer": str(bot_answer),
            "mode": str(mode)
        }

        with open(CHAT_HISTORY_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")

    except Exception as e:
        logger.error(f"event=chat_history_logging_failed | error={str(e)}")

CHAT_MODEL = "gpt-5.5"
VISION_MODEL = "gpt-5.5"
EMBEDDING_MODEL = "text-embedding-3-small"
RAG_THRESHOLD = 0.45


def load_embeddings():
    if not EMBEDDINGS_PATH.exists():
        logger.error(f"event=embeddings_file_missing | path={EMBEDDINGS_PATH}")
        raise FileNotFoundError(f"File not found: {EMBEDDINGS_PATH}")

    with open(EMBEDDINGS_PATH, "r", encoding="utf-8") as f:
        data = json.load(f)

    items = [row["item"] for row in data]
    embeddings = np.array([row["embedding"] for row in data], dtype=np.float32)
    logger.info(f"event=embeddings_loaded | items_count={len(items)}")
    return items, embeddings


knowledge_items, knowledge_embeddings = load_embeddings()


def prepare_text(item):
    return f"""
Type: {item.get("risk", "")}
Message: {item.get("text", "")}
Advice: {item.get("advice", "")}
""".strip()


def cosine_similarity(query_vector, matrix):
    query_norm_value = np.linalg.norm(query_vector)
    matrix_norm_value = np.linalg.norm(matrix, axis=1, keepdims=True)

    if query_norm_value == 0:
        return np.zeros(matrix.shape[0])

    matrix_norm_value[matrix_norm_value == 0] = 1
    query_norm = query_vector / query_norm_value
    matrix_norm = matrix / matrix_norm_value

    return np.dot(matrix_norm, query_norm)


@traceable(
    name="OpenAI Query Embedding",
    run_type="embedding",
    tags=["embedding", "openai", "rag"]
)
def embed_query(query):
    response = client.embeddings.create(
        model=EMBEDDING_MODEL,
        input=query[:6000]
    )
    return np.array(response.data[0].embedding, dtype=np.float32)


@traceable(
    name="RAG Retrieval - Embedding Search",
    run_type="retriever",
    tags=["rag", "retrieval", "embedding-search"]
)
def retrieve_by_embedding(query, top_k=3):
    query_embedding = embed_query(query)
    scores = cosine_similarity(query_embedding, knowledge_embeddings)
    top_indices = np.argsort(scores)[::-1][:top_k]

    results = []
    for idx in top_indices:
        item = knowledge_items[idx]
        results.append({
            "score": float(scores[idx]),
            "item": item,
            "text": prepare_text(item)
        })

    return results


def build_context(results):
    context_parts = []
    for i, result in enumerate(results, start=1):
        context_parts.append(
            f"""
Source {i}
Similarity Score: {round(result["score"], 4)}
{result["text"]}
"""
        )
    return "\n".join(context_parts)


def build_history_text(chat_history):
    if not chat_history:
        return "No previous conversation."

    history = ""
    for msg in chat_history[-10:]:
        role = msg.get("role", "")
        content = msg.get("content", "")

        if role == "user":
            history += f"User said: {content}\n"
        elif role == "assistant":
            history += f"Assistant answered: {content}\n"

    return history


def get_last_user_message(chat_history):
    for msg in reversed(chat_history):
        if msg.get("role") == "user":
            return msg.get("content", "")
    return ""


def is_follow_up(message):
    msg = message.lower()
    phrases = [
        "شو أعمل", "شو اعمل", "شو اسوي",
        "ماذا أفعل", "ماذا افعل",
        "هل اضغط", "هل أضغط", "هل افتح",
        "هل ارد", "هل أرد",
        "طيب", "يعني", "نفس الاشي",
        "is it safe", "what should i do",
        "should i click", "should i open"
    ]
    return any(p.lower() in msg for p in phrases)


def detect_user_intent(message):
    msg = message.strip().lower()

    educational_keywords = [
        "ما هو", "ماهي", "ما هي",
        "شو هو", "شو يعني",
        "ماذا يعني", "ماذا تعني",
        "ما معنى", "معنى",
        "اشرح", "اشرحلي",
        "عرف", "تعريف",
        "يعني ايش", "يعني شو",
        "what is", "explain", "define",
        "how does", "كيف يعمل",
        "كيف احمي", "كيف أحمي",
        "نصائح", "مفهوم", "concept"
    ]

    analysis_keywords = [
        "افحص", "حلل", "تحقق",
        "رسالة", "ايميل", "إيميل",
        "رابط", "link",
        "otp", "password", "كلمة المرور",
        "حسابك", "تم تعليق",
        "اضغط", "سجل دخول",
        "bank", "paypal",
        "phishing", "scam"
    ]

    if any(k in msg for k in educational_keywords):
        return "educational"

    if any(k in msg for k in analysis_keywords):
        return "analysis"

    return "educational"

def safe_json_parse(text):
    original_text = text
    try:
        text = text.strip()

        if text.startswith("```json"):
            text = text.replace("```json", "").replace("```", "").strip()
        elif text.startswith("```"):
            text = text.replace("```", "").strip()

        start = text.find("{")
        end = text.rfind("}")

        if start != -1 and end != -1:
            text = text[start:end + 1]

        return json.loads(text)

    except Exception as e:
        logger.error(f"event=json_parse_failed | error={str(e)}")
        return {
            "answer_type": "analysis",
            "classification": "Suspicious",
            "confidence": 0.45,
            "reason": "صار خطأ تقني أثناء قراءة رد النموذج، وليس حكمًا نهائيًا على الملف.",
            "red_flags": ["JSON parsing failed"],
            "advice": "أعد المحاولة مرة أخرى.",
            "raw_text": original_text
        }


def decode_text_bytes(data):
    encodings = ["utf-8", "utf-8-sig", "cp1256", "windows-1256", "iso-8859-6", "latin-1"]

    for enc in encodings:
        try:
            decoded = data.decode(enc, errors="strict")
            if decoded.strip():
                return decoded
        except Exception:
            pass

    return data.decode("utf-8", errors="ignore")


def clean_extracted_text(text):
    if not text:
        return ""

    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()[:15000]


def is_image_file(filename, content_type):
    name = filename.lower()
    return content_type.startswith("image/") or name.endswith((".png", ".jpg", ".jpeg", ".webp"))


def pdf_to_images(data, max_pages=3):
    image_urls = []

    try:
        import fitz

        doc = fitz.open(stream=data, filetype="pdf")

        for page_index in range(min(len(doc), max_pages)):
            page = doc.load_page(page_index)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            img_bytes = pix.tobytes("png")
            encoded = base64.b64encode(img_bytes).decode("utf-8")
            image_urls.append(f"data:image/png;base64,{encoded}")

        doc.close()

    except Exception:
        pass

    return image_urls


def docx_to_images(data, max_images=5):
    image_urls = []

    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            image_files = [
                name for name in z.namelist()
                if name.startswith("word/media/")
                and name.lower().endswith((".png", ".jpg", ".jpeg", ".webp"))
            ]

            for image_name in image_files[:max_images]:
                img_bytes = z.read(image_name)

                if image_name.lower().endswith(".png"):
                    mime = "image/png"
                elif image_name.lower().endswith(".webp"):
                    mime = "image/webp"
                else:
                    mime = "image/jpeg"

                encoded = base64.b64encode(img_bytes).decode("utf-8")
                image_urls.append(f"data:{mime};base64,{encoded}")

    except Exception:
        pass

    return image_urls


@traceable(
    name="Uploaded File Text Extraction",
    run_type="tool",
    tags=["file-analysis", "text-extraction"]
)
def extract_text_from_upload(filename, content_type, data):
    name = filename.lower()

    result = {
        "text": "",
        "readable": False,
        "method": "none",
        "note": "",
        "image_urls": []
    }

    try:
        if name.endswith((".txt", ".md", ".csv", ".json", ".log", ".html", ".xml")):
            text = clean_extracted_text(decode_text_bytes(data))
            result["text"] = text
            result["readable"] = bool(text.strip())
            result["method"] = "text_decode"
            return result

        if name.endswith(".pdf") or content_type == "application/pdf":
            try:
                from pypdf import PdfReader

                reader = PdfReader(io.BytesIO(data))
                text = ""

                for page in reader.pages[:10]:
                    text += (page.extract_text() or "") + "\n"

                text = clean_extracted_text(text)

                if text:
                    result["text"] = text
                    result["readable"] = True
                    result["method"] = "pypdf"
                    return result

            except Exception as e:
                result["note"] = f"pypdf failed: {str(e)}"

            image_urls = pdf_to_images(data, max_pages=3)
            result["image_urls"] = image_urls
            result["readable"] = bool(image_urls)
            result["method"] = "pdf_vision_fallback"
            result["note"] = "PDF analyzed visually because text extraction failed."
            return result

        if name.endswith(".docx") or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                import docx

                document = docx.Document(io.BytesIO(data))
                parts = []

                for p in document.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)

                for table in document.tables:
                    for row in table.rows:
                        cells = []
                        for cell in row.cells:
                            if cell.text.strip():
                                cells.append(cell.text.strip())
                        if cells:
                            parts.append(" | ".join(cells))

                text = clean_extracted_text("\n".join(parts))

                if text:
                    result["text"] = text
                    result["readable"] = True
                    result["method"] = "python_docx"
                    return result

            except Exception as e:
                result["note"] = f"docx failed: {str(e)}"

            image_urls = docx_to_images(data, max_images=5)
            result["image_urls"] = image_urls
            result["readable"] = bool(image_urls)
            result["method"] = "docx_vision_fallback"
            result["note"] = "DOCX analyzed visually because text extraction failed."
            return result

        fallback_text = clean_extracted_text(decode_text_bytes(data))

        if fallback_text and len(fallback_text.strip()) >= 10:
            result["text"] = fallback_text
            result["readable"] = True
            result["method"] = "fallback_decode"
            return result

        result["note"] = "Unsupported or binary file type."
        return result

    except Exception as e:
        result["note"] = f"General extraction error: {str(e)}"
        return result


@traceable(
    name="Rule Based Phishing Detection",
    run_type="tool",
    tags=["rules", "phishing-detection", "safety-check"]
)
def detect_phishing_rules(text):
    t = text.lower()
    red_flags = []

    has_password = any(k in t for k in [
        "password", "كلمة المرور", "كلمه المرور", "passcode"
    ])

    has_otp = any(k in t for k in [
        "otp", "one-time password", "verification code", "رمز التحقق", "كود التحقق"
    ])

    has_link = bool(re.search(r"https?://|www\.|bit\.ly|tinyurl|t\.co|wa\.me|forms\.gle", t))

    has_urgency = any(k in t for k in [
        "immediately", "urgent", "within 24 hours", "فوراً", "فورا", "عاجل", "خلال 24 ساعة"
    ])

    has_suspended = any(k in t for k in [
        "suspended", "account limited", "تم تعليق", "سيتم إغلاق", "حسابك موقوف"
    ])

    has_sensitive = any(k in t for k in [
        "bank account", "card number", "credit card", "cvv", "iban",
        "رقم البطاقة", "بياناتك البنكية", "رقم الهوية"
    ])

    has_login = any(k in t for k in [
        "login", "sign in", "verify your account",
        "سجل دخول", "تسجيل الدخول", "تحقق من حسابك"
    ])

    has_prize = any(k in t for k in [
        "you won", "winner", "lottery", "prize",
        "ربحت", "فزت", "جائزة", "مبروك لقد فزت"
    ])

    if has_password:
        red_flags.append("طلب كلمة المرور")
    if has_otp:
        red_flags.append("طلب رمز التحقق OTP")
    if has_link:
        red_flags.append("وجود رابط")
    if has_urgency:
        red_flags.append("استعجال أو تهديد زمني")
    if has_suspended:
        red_flags.append("ادعاء تعليق أو إغلاق الحساب")
    if has_sensitive:
        red_flags.append("طلب بيانات حساسة أو مالية")
    if has_login:
        red_flags.append("طلب تسجيل دخول")
    if has_prize:
        red_flags.append("وعد بجائزة أو ربح غير متوقع")

    strong_count = sum([
        has_password, has_otp, has_link, has_urgency,
        has_suspended, has_sensitive, has_login, has_prize
    ])

    if (has_password or has_otp or has_sensitive) and has_link:
        return {
            "answer_type": "analysis",
            "classification": "Phishing",
            "confidence": 0.92,
            "reason": "المحتوى يحتوي على مؤشرات تصيد قوية مثل طلب كلمة مرور أو رمز تحقق أو بيانات مالية مع وجود رابط.",
            "red_flags": red_flags,
            "advice": "لا تفتح الرابط ولا تدخل أي بيانات. استخدم الموقع أو التطبيق الرسمي فقط."
        }

    if has_link and has_login and (has_urgency or has_suspended):
        return {
            "answer_type": "analysis",
            "classification": "Phishing",
            "confidence": 0.88,
            "reason": "المحتوى يجمع بين رابط وتسجيل دخول واستعجال أو تهديد بالحساب.",
            "red_flags": red_flags,
            "advice": "لا تضغط الرابط. تحقق من الحساب من الموقع الرسمي."
        }

    if strong_count >= 4:
        return {
            "answer_type": "analysis",
            "classification": "Phishing",
            "confidence": 0.84,
            "reason": "المحتوى يحتوي على عدة مؤشرات خطيرة مجتمعة.",
            "red_flags": red_flags,
            "advice": "تعامل معه كتصيد ولا تشارك أي بيانات."
        }

    return None


def generate_safe_file_answer(filename):
    return {
        "answer_type": "analysis",
        "classification": "Safe",
        "confidence": 0.82,
        "reason": "تم تحليل محتوى الملف ولم تظهر مؤشرات تصيد واضحة مثل طلب كلمة مرور، OTP، بيانات بنكية، رابط تسجيل دخول مشبوه، تهديد، أو استعجال.",
        "red_flags": [],
        "advice": "يبدو الملف آمنًا من ناحية مؤشرات التصيد الظاهرة."
    }


def generate_unreadable_file_answer(filename, content_type, note):
    return {
        "answer_type": "analysis",
        "classification": "Needs Review",
        "confidence": 0.30,
        "reason": f"لم أستطع استخراج أو تحليل محتوى واضح من الملف {filename}. هذه حالة تقنية وليست حكمًا بأنه مشبوه.",
        "red_flags": [],
        "advice": "جرّب رفع نسخة أوضح من الملف أو صورة واضحة.",
        "technical_note": note,
        "content_type": content_type
    }


@traceable(
    name="Educational Cybersecurity Answer",
    run_type="llm",
    tags=["educational", "cybersecurity", "llm"]
)
def generate_educational_answer(req_message, history_text):
    system_prompt = """
You are a cybersecurity teacher.
Return ONLY valid JSON.

JSON schema:
{
  "answer_type": "educational",
  "title": "Arabic title",
  "explanation": "Detailed Arabic explanation",
  "examples": ["example 1", "example 2", "example 3"],
  "practical_tips": ["tip 1", "tip 2", "tip 3"],
  "short_summary": "short Arabic summary"
}
"""

    user_prompt = f"""
Previous conversation:
{history_text}

User question:
{req_message}

Answer in Arabic.
"""

    response = client.chat.completions.create(
        model=CHAT_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
    )

    raw = response.choices[0].message.content
    parsed = safe_json_parse(raw)

    # Ensure educational answers always have the full structure expected by Streamlit.
    if not isinstance(parsed, dict) or parsed.get("answer_type") != "educational":
        parsed = {
            "answer_type": "educational",
            "title": "شرح تعليمي",
            "explanation": raw,
            "examples": [],
            "practical_tips": [],
            "short_summary": raw[:300]
        }

    parsed.setdefault("answer_type", "educational")
    parsed.setdefault("title", "شرح تعليمي")
    parsed.setdefault("explanation", raw)
    parsed.setdefault("examples", [])
    parsed.setdefault("practical_tips", [])
    parsed.setdefault("short_summary", parsed.get("explanation", "")[:300])

    if not parsed.get("explanation"):
        parsed["explanation"] = raw

    return json.dumps(parsed, ensure_ascii=False)


@traceable(
    name="Uploaded Text File Analysis",
    run_type="llm",
    tags=["file-analysis", "text-file", "phishing-detection"]
)
def generate_text_file_analysis(message, filename, content_type, extracted_text, method):
    system_prompt = """
You are a careful cybersecurity phishing detector.
Analyze ONLY the provided file text.
Return ONLY valid JSON.

JSON schema:
{
  "answer_type": "analysis",
  "classification": "Safe" | "Suspicious" | "Phishing",
  "confidence": number between 0 and 1,
  "reason": "Arabic explanation",
  "red_flags": ["Arabic red flag 1", "Arabic red flag 2"],
  "advice": "Arabic practical advice"
}

STRICT RULES:
- Do NOT classify every uploaded file as suspicious.
- If the file is a normal CV, cover letter, homework, lecture, report, or assignment with no risky request, classify Safe.
- Phishing only if there is clear evidence: password, OTP, bank/card info, fake login, suspicious link, urgent account suspension, impersonation, or malware-like instruction.
- Suspicious only if there are real weak indicators.
"""

    user_prompt = f"""
User message:
{message}

File name:
{filename}

File type:
{content_type}

Extraction method:
{method}

Extracted text:
{extracted_text}

Return only valid JSON.
"""

    response = client.chat.completions.create(
        model=CHAT_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
    )

    return response.choices[0].message.content


@traceable(
    name="Uploaded Image or PDF Vision Analysis",
    run_type="llm",
    tags=["vision-analysis", "file-analysis", "image-pdf"]
)
def generate_vision_analysis(message, filename, content_type, image_urls):
    system_prompt = """
You are a careful cybersecurity phishing detector with vision ability.
Analyze ONLY the provided image/PDF/DOCX pages.
Return ONLY valid JSON.

JSON schema:
{
  "answer_type": "analysis",
  "classification": "Safe" | "Suspicious" | "Phishing",
  "confidence": number between 0 and 1,
  "reason": "Arabic explanation",
  "red_flags": ["Arabic red flag 1", "Arabic red flag 2"],
  "advice": "Arabic practical advice"
}

STRICT RULES:
- Do NOT classify every uploaded file as suspicious.
- If it is a normal CV, cover letter, homework, assignment, report, lecture slide, invoice, or certificate with no risky request, classify Safe.
- Phishing only if there is clear evidence: password request, OTP request, bank/card request, fake login, suspicious link, urgent account suspension, impersonation, or malware-like instruction.
- Suspicious only if there are real weak indicators.
"""

    content = [{
        "type": "text",
        "text": f"""
User message:
{message}

File name:
{filename}

File type:
{content_type}

Analyze visible content only.
Return only valid JSON.
"""
    }]

    for url in image_urls:
        content.append({
            "type": "image_url",
            "image_url": {"url": url}
        })

    response = client.chat.completions.create(
        model=VISION_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": content}
        ]
    )

    return response.choices[0].message.content


@traceable(
    name="RAG Answer Generation",
    run_type="llm",
    tags=["rag", "llm-generation", "phishing-analysis"]
)
def generate_answer_with_rag(req_message, history_text, follow_up, retrieval_query, context):
    system_prompt = """
You are a cybersecurity assistant specialized in scam and phishing detection.
Return ONLY valid JSON.

JSON schema:
{
  "answer_type": "analysis",
  "classification": "Safe" | "Suspicious" | "Phishing",
  "confidence": number between 0 and 1,
  "reason": "Arabic explanation",
  "red_flags": ["Arabic red flag 1", "Arabic red flag 2"],
  "advice": "Arabic practical advice"
}
"""

    user_prompt = f"""
Previous conversation:
{history_text}

Current user message:
{req_message}

Is follow-up:
{follow_up}

Message used for retrieval:
{retrieval_query}

Retrieved context:
{context}

Return only valid JSON.
"""

    response = client.chat.completions.create(
        model=CHAT_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
    )

    return response.choices[0].message.content


@traceable(
    name="LLM Fallback Answer Generation",
    run_type="llm",
    tags=["llm-fallback", "phishing-analysis"]
)
def generate_answer_with_llm(req_message, history_text, follow_up, retrieval_query):
    system_prompt = """
You are a cybersecurity assistant specialized in scam and phishing detection.
Return ONLY valid JSON.

JSON schema:
{
  "answer_type": "analysis",
  "classification": "Safe" | "Suspicious" | "Phishing",
  "confidence": number between 0 and 1,
  "reason": "Arabic explanation",
  "red_flags": ["Arabic red flag 1", "Arabic red flag 2"],
  "advice": "Arabic practical advice"
}
"""

    user_prompt = f"""
Previous conversation:
{history_text}

Current user message:
{req_message}

Is follow-up:
{follow_up}

Message/file being analyzed:
{retrieval_query}

Return only valid JSON.
"""

    response = client.chat.completions.create(
        model=CHAT_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
    )

    return response.choices[0].message.content


def adjust_confidence(answer, retrieved, mode):
    classification = answer.get("classification", "Suspicious")
    confidence = float(answer.get("confidence", 0.5))

    top_score = retrieved[0]["score"] if retrieved else 0
    avg_score = sum(r["score"] for r in retrieved) / len(retrieved) if retrieved else 0

    if classification == "Phishing":
        confidence = min(max(confidence, 0.75), 0.95)
    elif classification == "Suspicious":
        confidence = min(max(confidence, 0.40), 0.75)
    elif classification == "Safe":
        confidence = min(max(confidence, 0.65), 0.90)
    elif classification == "Needs Review":
        confidence = min(max(confidence, 0.20), 0.45)

    answer["confidence"] = round(confidence, 2)
    answer["retrieval_top_score"] = round(top_score, 4)
    answer["retrieval_avg_score"] = round(avg_score, 4)
    answer["answer_mode"] = mode

    return answer


@app.get("/")
def home():
    return {
        "status": "running",
        "version": "LANGSMITH_MONITORING_ENABLED",
        "system_type": "hybrid-rag-llm-file-image-pdf-docx-vision",
        "chat_model": CHAT_MODEL,
        "vision_model": VISION_MODEL,
        "embedding_model": EMBEDDING_MODEL,
        "knowledge_items": len(knowledge_items),
        "rag_threshold": RAG_THRESHOLD,
        "langsmith_project": os.getenv("LANGSMITH_PROJECT", "LLMProject")
    }


@app.post("/chat")
@traceable(
    name="Scam Phishing RAG Chatbot",
    run_type="chain",
    tags=["scam-detection", "phishing", "rag-chatbot"]
)
async def chat(
    message: str = Form(""),
    chat_history: str = Form("[]"),
    files: Optional[List[UploadFile]] = File(None)
):
    start_time = time.time()
    request_id = str(uuid.uuid4())[:8]

    logger.info(
        f"request_id={request_id} | event=chat_request_started | "
        f"message_length={len(message)} | has_files={files is not None}"
    )

    try:
        parsed_history = json.loads(chat_history)
    except Exception as e:
        logger.warning(f"request_id={request_id} | event=chat_history_parse_failed | error={str(e)}")
        parsed_history = []

    history_text = build_history_text(parsed_history)

    uploaded_files_info = []
    file_items = []

    if files:
        for uploaded in files:
            data = await uploaded.read()
            filename = uploaded.filename or "unknown_file"
            content_type = uploaded.content_type or "application/octet-stream"

            logger.info(
                f"request_id={request_id} | event=file_uploaded | "
                f"filename={filename} | content_type={content_type} | size_bytes={len(data)}"
            )

            uploaded_files_info.append({
                "filename": filename,
                "content_type": content_type,
                "size_bytes": len(data)
            })

            if is_image_file(filename, content_type):
                encoded = base64.b64encode(data).decode("utf-8")
                image_url = f"data:{content_type};base64,{encoded}"

                file_items.append({
                    "filename": filename,
                    "content_type": content_type,
                    "text": "",
                    "image_urls": [image_url],
                    "method": "image_upload",
                    "note": ""
                })

            else:
                extraction = extract_text_from_upload(filename, content_type, data)

                logger.info(
                    f"request_id={request_id} | event=file_extracted | "
                    f"filename={filename} | method={extraction.get('method')} | "
                    f"readable={extraction.get('readable')} | "
                    f"text_length={len(extraction.get('text', ''))} | "
                    f"image_pages={len(extraction.get('image_urls', []))}"
                )

                file_items.append({
                    "filename": filename,
                    "content_type": content_type,
                    "text": extraction.get("text", ""),
                    "image_urls": extraction.get("image_urls", []),
                    "method": extraction.get("method", "none"),
                    "note": extraction.get("note", "")
                })

    if uploaded_files_info:
        file_results = []

        for item in file_items:
            filename = item["filename"]
            content_type = item["content_type"]
            extracted_text = item.get("text", "").strip()
            image_urls = item.get("image_urls", [])
            method = item.get("method", "none")
            note = item.get("note", "")

            if extracted_text:
                rule_answer = detect_phishing_rules(extracted_text)

                if rule_answer:
                    parsed_answer = rule_answer
                    raw_answer = json.dumps(rule_answer, ensure_ascii=False)
                else:
                    raw_answer = generate_text_file_analysis(
                        message=message,
                        filename=filename,
                        content_type=content_type,
                        extracted_text=extracted_text,
                        method=method
                    )
                    parsed_answer = safe_json_parse(raw_answer)

                    if parsed_answer.get("classification") == "Suspicious" and not parsed_answer.get("red_flags"):
                        parsed_answer = generate_safe_file_answer(filename)
                        raw_answer = json.dumps(parsed_answer, ensure_ascii=False)

                parsed_answer = adjust_confidence(parsed_answer, [], "FILE_TEXT_ANALYSIS")
                mode = "FILE_TEXT_ANALYSIS"
                preview = extracted_text[:700]

            elif image_urls:
                raw_answer = generate_vision_analysis(
                    message=message,
                    filename=filename,
                    content_type=content_type,
                    image_urls=image_urls
                )
                parsed_answer = safe_json_parse(raw_answer)
                parsed_answer = adjust_confidence(parsed_answer, [], "VISION_FILE_ANALYSIS")
                mode = "VISION_FILE_ANALYSIS"
                preview = "[Analyzed visually]"

            else:
                parsed_answer = generate_unreadable_file_answer(
                    filename=filename,
                    content_type=content_type,
                    note=note
                )
                parsed_answer = adjust_confidence(parsed_answer, [], "UNREADABLE_FILE")
                raw_answer = json.dumps(parsed_answer, ensure_ascii=False)
                mode = "UNREADABLE_FILE"
                preview = ""

            logger.info(
                f"request_id={request_id} | event=file_analysis_result | "
                f"filename={filename} | mode={mode} | "
                f"classification={parsed_answer.get('classification')} | "
                f"confidence={parsed_answer.get('confidence')}"
            )

            file_results.append({
                "filename": filename,
                "content_type": content_type,
                "mode": mode,
                "answer": parsed_answer,
                "raw_answer": raw_answer,
                "extracted_preview": preview,
                "extraction_method": method,
                "extraction_note": note
            })

        latency = round(time.time() - start_time, 2)

        logger.info(
            f"request_id={request_id} | event=file_analysis_completed | "
            f"latency={latency}s | files_count={len(file_results)}"
        )

        return {
            "mode": "MULTI_FILE_IMAGE_ANALYSIS",
            "system_type": "hybrid-rag-llm-file-image-pdf-docx-vision",
            "intent": "file_or_image_analysis",
            "latency": latency,
            "file_results": file_results,
            "uploaded_files": uploaded_files_info,
            "retrieved_sources": []
        }

    intent = detect_user_intent(message)

    logger.info(f"request_id={request_id} | event=intent_detected | intent={intent}")

    if intent == "educational":
        raw_answer = generate_educational_answer(message, history_text)
        parsed_answer = safe_json_parse(raw_answer)
        latency = round(time.time() - start_time, 2)

        logger.info(
            f"request_id={request_id} | event=educational_answer_completed | "
            f"latency={latency}s | intent={intent}"
        )

        return {
            "mode": "LLM_EDUCATIONAL",
            "system_type": "hybrid-rag-llm-file-image-pdf-docx-vision",
            "intent": intent,
            "latency": latency,
            "answer": parsed_answer,
            "raw_answer": raw_answer,
            "retrieved_sources": []
        }

    follow_up = is_follow_up(message)
    last_user_message = get_last_user_message(parsed_history)

    if follow_up and last_user_message:
        retrieval_query = last_user_message + "\n" + message
    else:
        retrieval_query = message

    retrieved = retrieve_by_embedding(retrieval_query, top_k=3)
    context = build_context(retrieved)
    top_score = retrieved[0]["score"] if retrieved else 0

    logger.info(
        f"request_id={request_id} | event=retrieval_completed | "
        f"top_score={round(top_score, 4)} | threshold={RAG_THRESHOLD} | follow_up={follow_up}"
    )

    if top_score >= RAG_THRESHOLD:
        mode = "RAG_ANALYSIS"
        raw_answer = generate_answer_with_rag(
            req_message=message,
            history_text=history_text,
            follow_up=follow_up,
            retrieval_query=retrieval_query,
            context=context
        )
    else:
        mode = "LLM_ANALYSIS"
        raw_answer = generate_answer_with_llm(
            req_message=message,
            history_text=history_text,
            follow_up=follow_up,
            retrieval_query=retrieval_query
        )

    parsed_answer = safe_json_parse(raw_answer)
    parsed_answer = adjust_confidence(parsed_answer, retrieved, mode)

    latency = round(time.time() - start_time, 2)

    logger.info(
        f"request_id={request_id} | event=chat_answer_completed | "
        f"mode={mode} | intent={intent} | "
        f"classification={parsed_answer.get('classification')} | "
        f"confidence={parsed_answer.get('confidence')} | "
        f"top_score={round(top_score, 4)} | latency={latency}s"
    )

    # =========================
    # SAVE CHAT FOR EVALUATION
    # =========================
    try:
        save_chat_history(
            user_message=message,
            bot_answer=json.dumps(parsed_answer, ensure_ascii=False),
            mode=mode
        )
    except Exception as e:
        logger.error(f"event=chat_history_save_failed | error={str(e)}")

    return {
        "mode": mode,
        "system_type": "hybrid-rag-llm-file-image-pdf-docx-vision",
        "intent": intent,
        "is_follow_up": follow_up,
        "latency": latency,
        "rag_threshold": RAG_THRESHOLD,
        "answer": parsed_answer,
        "raw_answer": raw_answer,
        "retrieval_query": retrieval_query,
        "retrieved_sources": [
            {
                "score": round(result["score"], 4),
                "risk": result["item"].get("risk"),
                "text": result["item"].get("text", "")[:300],
                "advice": result["item"].get("advice")
            }
            for result in retrieved
        ]
    }