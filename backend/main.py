import json
import time
import base64
import re
import io
import zipfile
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional

import numpy as np
from fastapi import FastAPI, UploadFile, File, Form
from openai import OpenAI
from monitoring import setup_monitoring
import os
from datetime import datetime

# =========================
# LangSmith Setup
# =========================
from dotenv import load_dotenv
from langsmith import traceable

load_dotenv()
os.environ.setdefault("LANGSMITH_TRACING", "true")
os.environ.setdefault("LANGSMITH_PROJECT", "LLMProject")
os.environ.setdefault("LANGSMITH_ENDPOINT", "https://api.smith.langchain.com")


app = FastAPI(title="Scam Phishing Hybrid API")

client = OpenAI()

BASE_DIR = Path(__file__).resolve().parent
EMBEDDINGS_PATH = BASE_DIR / "data" / "phishing_embeddings.json"

logger = setup_monitoring(app, BASE_DIR)

# =========================
# Chat History Logging (FOR EVALUATION)
# =========================

CHAT_HISTORY_PATH = BASE_DIR / "logs" / "chat_history.jsonl"

def save_chat_history(user_message, bot_answer, mode="unknown"):
    try:
        os.makedirs(CHAT_HISTORY_PATH.parent, exist_ok=True)

        record = {
            "timestamp": datetime.now().isoformat(),
            "user_message": str(user_message),
            "bot_answer": str(bot_answer),
            "mode": str(mode)
        }

        with open(CHAT_HISTORY_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")

    except Exception as e:
        logger.error(f"event=chat_history_logging_failed | error={str(e)}")


# =========================
# A/B Testing Logging (FOR ADMIN DASHBOARD)
# =========================

AB_TEST_LOG_PATH = BASE_DIR / "logs" / "ab_testing_results.jsonl"

def save_ab_test_result(
    user_message,
    answer,
    mode="unknown",
    model_used="unknown",
    ab_variant="unknown",
    latency=0,
    top_score=0,
    intent="unknown",
    request_id="",
    filename=""
):
    """
    Saves A/B testing metadata for the admin dashboard only.
    The normal user chat UI should not display this information.
    """
    try:
        os.makedirs(AB_TEST_LOG_PATH.parent, exist_ok=True)

        if not isinstance(answer, dict):
            answer = {"raw_answer": str(answer)}

        record = {
            "timestamp": datetime.now().isoformat(),
            "request_id": str(request_id),
            "user_message": str(user_message),
            "filename": str(filename),
            "answer": answer,
            "answer_text": json.dumps(answer, ensure_ascii=False),
            "mode": str(mode),
            "intent": str(intent),
            "model_used": str(model_used),
            "ab_variant": str(ab_variant),
            "latency": float(latency or 0),
            "top_score": float(top_score or 0),
            "classification": str(answer.get("classification", "")),
            "confidence": float(answer.get("confidence", 0) or 0),
            "reason": str(answer.get("reason", "") or answer.get("explanation", "")),
            "advice": str(answer.get("advice", "") or answer.get("short_summary", "")),
        }

        with open(AB_TEST_LOG_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")

    except Exception as e:
        logger.error(f"event=ab_test_logging_failed | error={str(e)}")


# =========================
# Static File Risk Logging
# =========================

STATIC_FILE_RISK_LOG_PATH = BASE_DIR / "logs" / "static_file_risks.jsonl"

def save_static_file_risk_event(filename, content_type, static_risk, request_id="", user_message=""):
    """
    Saves static file/image risk events so the dashboard can update when files/images are analyzed.
    This does not execute the uploaded file; it only logs safe metadata and detected static indicators.
    """
    try:
        if not isinstance(static_risk, dict):
            return

        severity = int(static_risk.get("file_static_severity", 0) or 0)
        red_flags = static_risk.get("red_flags", []) or []

        # Log only meaningful static file/image signals.
        if severity <= 0 and not red_flags:
            return

        os.makedirs(STATIC_FILE_RISK_LOG_PATH.parent, exist_ok=True)

        record = {
            "timestamp": datetime.now().isoformat(),
            "request_id": str(request_id),
            "filename": str(filename),
            "content_type": str(content_type),
            "classification": str(static_risk.get("classification", "Suspicious")),
            "confidence": float(static_risk.get("confidence", 0.0) or 0.0),
            "file_static_severity": severity,
            "red_flags": red_flags,
            "reason": str(static_risk.get("reason", "")),
            "advice": str(static_risk.get("advice", "")),
            "user_message": str(user_message),
        }

        with open(STATIC_FILE_RISK_LOG_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")

    except Exception as e:
        logger.error(f"event=static_file_risk_logging_failed | error={str(e)}")

CHAT_MODEL = "gpt-5.5"
VISION_MODEL = "gpt-5.5"
EMBEDDING_MODEL = "text-embedding-3-small"
RAG_THRESHOLD = 0.45

# =========================
# A/B Testing Models
# =========================
import random
import contextvars

MODEL_A_NAME = "gpt-5.5"
MODEL_B_NAME = "microsoft/Phi-3.5-mini-instruct"

CURRENT_AB_MODEL = contextvars.ContextVar("CURRENT_AB_MODEL", default=MODEL_A_NAME)
LAST_USED_MODEL = contextvars.ContextVar("LAST_USED_MODEL", default=MODEL_A_NAME)
LAST_AB_VARIANT = contextvars.ContextVar("LAST_AB_VARIANT", default="A")

_phi_pipeline = None

def choose_ab_model(ab_model="auto"):
    """
    auto = random A/B testing
    gpt = force GPT
    phi = force Phi
    """
    ab_model = (ab_model or "auto").lower().strip()

    if ab_model == "gpt":
        variant = "A"
        model = MODEL_A_NAME
    elif ab_model == "phi":
        variant = "B"
        model = MODEL_B_NAME
    else:
        if random.random() < 0.5:
            variant = "A"
            model = MODEL_A_NAME
        else:
            variant = "B"
            model = MODEL_B_NAME

    CURRENT_AB_MODEL.set(model)
    LAST_USED_MODEL.set(model)
    LAST_AB_VARIANT.set(variant)
    return model, variant

def load_phi_pipeline():
    global _phi_pipeline

    if _phi_pipeline is None:
        from transformers import pipeline
        _phi_pipeline = pipeline(
            "text-generation",
            model=MODEL_B_NAME,
            device_map="auto",
            max_new_tokens=500
        )

    return _phi_pipeline

def call_ab_chat_model(system_prompt, user_prompt, allow_phi=True):
    """
    Unified model caller for A/B testing.
    GPT is Model A, Phi is Model B.
    If Phi fails, the system safely falls back to GPT so the project does not crash.
    """
    selected_model = CURRENT_AB_MODEL.get()

    if selected_model == MODEL_B_NAME and allow_phi:
        try:
            pipe = load_phi_pipeline()

            prompt = f"""
System:
{system_prompt}

User:
{user_prompt}

Return ONLY valid JSON.
"""
            output = pipe(prompt, do_sample=False)[0]["generated_text"]

            if "Return ONLY valid JSON." in output:
                output = output.split("Return ONLY valid JSON.")[-1].strip()

            LAST_USED_MODEL.set(MODEL_B_NAME)
            LAST_AB_VARIANT.set("B")
            return output

        except Exception as e:
            logger.error(f"event=phi_model_failed_fallback_to_gpt | error={str(e)}")

    response = client.chat.completions.create(
        model=MODEL_A_NAME,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
    )

    LAST_USED_MODEL.set(MODEL_A_NAME)
    LAST_AB_VARIANT.set("A")
    return response.choices[0].message.content

def add_ab_metadata(answer):
    if isinstance(answer, dict):
        answer["model_used"] = LAST_USED_MODEL.get()
        answer["ab_variant"] = LAST_AB_VARIANT.get()
        answer["ab_test_enabled"] = True
    return answer


def load_embeddings():
    if not EMBEDDINGS_PATH.exists():
        logger.error(f"event=embeddings_file_missing | path={EMBEDDINGS_PATH}")
        raise FileNotFoundError(f"File not found: {EMBEDDINGS_PATH}")

    with open(EMBEDDINGS_PATH, "r", encoding="utf-8") as f:
        data = json.load(f)

    items = [row["item"] for row in data]
    embeddings = np.array([row["embedding"] for row in data], dtype=np.float32)
    logger.info(f"event=embeddings_loaded | items_count={len(items)}")
    return items, embeddings


knowledge_items, knowledge_embeddings = load_embeddings()


def prepare_text(item):
    return f"""
Type: {item.get("risk", "")}
Message: {item.get("text", "")}
Advice: {item.get("advice", "")}
""".strip()


def cosine_similarity(query_vector, matrix):
    query_norm_value = np.linalg.norm(query_vector)
    matrix_norm_value = np.linalg.norm(matrix, axis=1, keepdims=True)

    if query_norm_value == 0:
        return np.zeros(matrix.shape[0])

    matrix_norm_value[matrix_norm_value == 0] = 1
    query_norm = query_vector / query_norm_value
    matrix_norm = matrix / matrix_norm_value

    return np.dot(matrix_norm, query_norm)


@traceable(
    name="OpenAI Query Embedding",
    run_type="embedding",
    tags=["embedding", "openai", "rag"]
)
def embed_query(query):
    response = client.embeddings.create(
        model=EMBEDDING_MODEL,
        input=query[:6000]
    )
    return np.array(response.data[0].embedding, dtype=np.float32)


@traceable(
    name="RAG Retrieval - Embedding Search",
    run_type="retriever",
    tags=["rag", "retrieval", "embedding-search"]
)
def retrieve_by_embedding(query, top_k=3):
    query_embedding = embed_query(query)
    scores = cosine_similarity(query_embedding, knowledge_embeddings)
    top_indices = np.argsort(scores)[::-1][:top_k]

    results = []
    for idx in top_indices:
        item = knowledge_items[idx]
        results.append({
            "score": float(scores[idx]),
            "item": item,
            "text": prepare_text(item)
        })

    return results


def build_context(results):
    context_parts = []
    for i, result in enumerate(results, start=1):
        context_parts.append(
            f"""
Source {i}
Similarity Score: {round(result["score"], 4)}
{result["text"]}
"""
        )
    return "\n".join(context_parts)


def build_history_text(chat_history):
    if not chat_history:
        return "No previous conversation."

    history = ""
    for msg in chat_history[-10:]:
        role = msg.get("role", "")
        content = msg.get("content", "")

        if role == "user":
            history += f"User said: {content}\n"
        elif role == "assistant":
            history += f"Assistant answered: {content}\n"

    return history


def get_last_user_message(chat_history):
    for msg in reversed(chat_history):
        if msg.get("role") == "user":
            return msg.get("content", "")
    return ""


def is_follow_up(message):
    msg = message.lower()
    phrases = [
        "شو أعمل", "شو اعمل", "شو اسوي",
        "ماذا أفعل", "ماذا افعل",
        "هل اضغط", "هل أضغط", "هل افتح",
        "هل ارد", "هل أرد",
        "طيب", "يعني", "نفس الاشي",
        "is it safe", "what should i do",
        "should i click", "should i open"
    ]
    return any(p.lower() in msg for p in phrases)


def detect_user_intent(message):
    msg = message.strip().lower()

    educational_keywords = [
        "ما هو", "ماهي", "ما هي",
        "شو هو", "شو يعني",
        "ماذا يعني", "ماذا تعني",
        "ما معنى", "معنى",
        "اشرح", "اشرحلي",
        "عرف", "تعريف",
        "يعني ايش", "يعني شو",
        "what is", "explain", "define",
        "how does", "كيف يعمل",
        "كيف احمي", "كيف أحمي",
        "نصائح", "مفهوم", "concept"
    ]

    analysis_keywords = [
        "افحص", "حلل", "تحقق",
        "رسالة", "ايميل", "إيميل",
        "رابط", "link",
        "otp", "password", "كلمة المرور",
        "حسابك", "تم تعليق",
        "اضغط", "سجل دخول",
        "bank", "paypal",
        "phishing", "scam"
    ]

    if any(k in msg for k in educational_keywords):
        return "educational"

    if any(k in msg for k in analysis_keywords):
        return "analysis"

    return "educational"



def detect_response_language(text):
    """
    Detects the dominant language of the user's current message.
    This does not affect RAG, evaluation, dashboard, or monitoring.
    It only tells the LLM what language to use in the final JSON values.
    """
    text = text or ""
    arabic_chars = len(re.findall(r"[\u0600-\u06FF]", text))
    english_chars = len(re.findall(r"[A-Za-z]", text))

    if arabic_chars > english_chars:
        return "Arabic"
    if english_chars > arabic_chars:
        return "English"
    return "the same language as the user's message"


def language_instruction(text):
    lang = detect_response_language(text)
    return f"""
IMPORTANT LANGUAGE RULE:
- The current user message language is: {lang}.
- Use {lang} for EVERY human-readable JSON value.
- This includes: title, explanation, reason, red_flags, advice, examples, practical_tips, short_summary.
- Do NOT mix Arabic and English inside reason, red_flags, or advice.
- Keep JSON keys in English exactly as required by the schema.
- Keep classification values exactly one of: Safe, Suspicious, Phishing, Needs Review.
""".strip()

def safe_json_parse(text):
    original_text = text
    try:
        text = text.strip()

        if text.startswith("```json"):
            text = text.replace("```json", "").replace("```", "").strip()
        elif text.startswith("```"):
            text = text.replace("```", "").strip()

        start = text.find("{")
        end = text.rfind("}")

        if start != -1 and end != -1:
            text = text[start:end + 1]

        return json.loads(text)

    except Exception as e:
        logger.error(f"event=json_parse_failed | error={str(e)}")
        return {
            "answer_type": "analysis",
            "classification": "Suspicious",
            "confidence": 0.45,
            "reason": "صار خطأ تقني أثناء قراءة رد النموذج، وليس حكمًا نهائيًا على الملف.",
            "red_flags": ["JSON parsing failed"],
            "advice": "أعد المحاولة مرة أخرى.",
            "raw_text": original_text
        }


def decode_text_bytes(data):
    encodings = ["utf-8", "utf-8-sig", "cp1256", "windows-1256", "iso-8859-6", "latin-1"]

    for enc in encodings:
        try:
            decoded = data.decode(enc, errors="strict")
            if decoded.strip():
                return decoded
        except Exception:
            pass

    return data.decode("utf-8", errors="ignore")


def clean_extracted_text(text):
    if not text:
        return ""

    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()[:15000]


def is_image_file(filename, content_type):
    name = filename.lower()
    return content_type.startswith("image/") or name.endswith((".png", ".jpg", ".jpeg", ".webp"))


def pdf_to_images(data, max_pages=3):
    image_urls = []

    try:
        import fitz

        doc = fitz.open(stream=data, filetype="pdf")

        for page_index in range(min(len(doc), max_pages)):
            page = doc.load_page(page_index)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            img_bytes = pix.tobytes("png")
            encoded = base64.b64encode(img_bytes).decode("utf-8")
            image_urls.append(f"data:image/png;base64,{encoded}")

        doc.close()

    except Exception:
        pass

    return image_urls


def docx_to_images(data, max_images=5):
    image_urls = []

    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            image_files = [
                name for name in z.namelist()
                if name.startswith("word/media/")
                and name.lower().endswith((".png", ".jpg", ".jpeg", ".webp"))
            ]

            for image_name in image_files[:max_images]:
                img_bytes = z.read(image_name)

                if image_name.lower().endswith(".png"):
                    mime = "image/png"
                elif image_name.lower().endswith(".webp"):
                    mime = "image/webp"
                else:
                    mime = "image/jpeg"

                encoded = base64.b64encode(img_bytes).decode("utf-8")
                image_urls.append(f"data:{mime};base64,{encoded}")

    except Exception:
        pass

    return image_urls


@traceable(
    name="Uploaded File Text Extraction",
    run_type="tool",
    tags=["file-analysis", "text-extraction"]
)
def extract_text_from_upload(filename, content_type, data):
    name = filename.lower()

    result = {
        "text": "",
        "readable": False,
        "method": "none",
        "note": "",
        "image_urls": []
    }

    try:
        if name.endswith((".txt", ".md", ".csv", ".json", ".log", ".html", ".xml")):
            text = clean_extracted_text(decode_text_bytes(data))
            result["text"] = text
            result["readable"] = bool(text.strip())
            result["method"] = "text_decode"
            return result

        if name.endswith(".pdf") or content_type == "application/pdf":
            try:
                from pypdf import PdfReader

                reader = PdfReader(io.BytesIO(data))
                text = ""

                for page in reader.pages[:10]:
                    text += (page.extract_text() or "") + "\n"

                text = clean_extracted_text(text)

                if text:
                    result["text"] = text
                    result["readable"] = True
                    result["method"] = "pypdf"
                    return result

            except Exception as e:
                result["note"] = f"pypdf failed: {str(e)}"

            image_urls = pdf_to_images(data, max_pages=3)
            result["image_urls"] = image_urls
            result["readable"] = bool(image_urls)
            result["method"] = "pdf_vision_fallback"
            result["note"] = "PDF analyzed visually because text extraction failed."
            return result

        if name.endswith(".docx") or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                import docx

                document = docx.Document(io.BytesIO(data))
                parts = []

                for p in document.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)

                for table in document.tables:
                    for row in table.rows:
                        cells = []
                        for cell in row.cells:
                            if cell.text.strip():
                                cells.append(cell.text.strip())
                        if cells:
                            parts.append(" | ".join(cells))

                text = clean_extracted_text("\n".join(parts))

                if text:
                    result["text"] = text
                    result["readable"] = True
                    result["method"] = "python_docx"
                    return result

            except Exception as e:
                result["note"] = f"docx failed: {str(e)}"

            image_urls = docx_to_images(data, max_images=5)
            result["image_urls"] = image_urls
            result["readable"] = bool(image_urls)
            result["method"] = "docx_vision_fallback"
            result["note"] = "DOCX analyzed visually because text extraction failed."
            return result

        fallback_text = clean_extracted_text(decode_text_bytes(data))

        if fallback_text and len(fallback_text.strip()) >= 10:
            result["text"] = fallback_text
            result["readable"] = True
            result["method"] = "fallback_decode"
            return result

        result["note"] = "Unsupported or binary file type."
        return result

    except Exception as e:
        result["note"] = f"General extraction error: {str(e)}"
        return result


@traceable(
    name="Rule Based Phishing Detection",
    run_type="tool",
    tags=["rules", "phishing-detection", "safety-check"]
)
def detect_phishing_rules(text):
    t = text.lower()
    red_flags = []

    has_password = any(k in t for k in [
        "password", "كلمة المرور", "كلمه المرور", "passcode"
    ])

    has_otp = any(k in t for k in [
        "otp", "one-time password", "verification code", "رمز التحقق", "كود التحقق"
    ])

    has_link = bool(re.search(r"https?://|www\.|bit\.ly|tinyurl|t\.co|wa\.me|forms\.gle", t))

    has_urgency = any(k in t for k in [
        "immediately", "urgent", "within 24 hours", "فوراً", "فورا", "عاجل", "خلال 24 ساعة"
    ])

    has_suspended = any(k in t for k in [
        "suspended", "account limited", "تم تعليق", "سيتم إغلاق", "حسابك موقوف"
    ])

    has_sensitive = any(k in t for k in [
        "bank account", "card number", "credit card", "cvv", "iban",
        "رقم البطاقة", "بياناتك البنكية", "رقم الهوية"
    ])

    has_login = any(k in t for k in [
        "login", "sign in", "verify your account",
        "سجل دخول", "تسجيل الدخول", "تحقق من حسابك"
    ])

    has_prize = any(k in t for k in [
        "you won", "winner", "lottery", "prize",
        "ربحت", "فزت", "جائزة", "مبروك لقد فزت"
    ])

    if has_password:
        red_flags.append("طلب كلمة المرور")
    if has_otp:
        red_flags.append("طلب رمز التحقق OTP")
    if has_link:
        red_flags.append("وجود رابط")
    if has_urgency:
        red_flags.append("استعجال أو تهديد زمني")
    if has_suspended:
        red_flags.append("ادعاء تعليق أو إغلاق الحساب")
    if has_sensitive:
        red_flags.append("طلب بيانات حساسة أو مالية")
    if has_login:
        red_flags.append("طلب تسجيل دخول")
    if has_prize:
        red_flags.append("وعد بجائزة أو ربح غير متوقع")

    strong_count = sum([
        has_password, has_otp, has_link, has_urgency,
        has_suspended, has_sensitive, has_login, has_prize
    ])

    if (has_password or has_otp or has_sensitive) and has_link:
        return {
            "answer_type": "analysis",
            "classification": "Phishing",
            "confidence": 0.92,
            "reason": "المحتوى يحتوي على مؤشرات تصيد قوية مثل طلب كلمة مرور أو رمز تحقق أو بيانات مالية مع وجود رابط.",
            "red_flags": red_flags,
            "advice": "لا تفتح الرابط ولا تدخل أي بيانات. استخدم الموقع أو التطبيق الرسمي فقط."
        }

    if has_link and has_login and (has_urgency or has_suspended):
        return {
            "answer_type": "analysis",
            "classification": "Phishing",
            "confidence": 0.88,
            "reason": "المحتوى يجمع بين رابط وتسجيل دخول واستعجال أو تهديد بالحساب.",
            "red_flags": red_flags,
            "advice": "لا تضغط الرابط. تحقق من الحساب من الموقع الرسمي."
        }

    if strong_count >= 4:
        return {
            "answer_type": "analysis",
            "classification": "Phishing",
            "confidence": 0.84,
            "reason": "المحتوى يحتوي على عدة مؤشرات خطيرة مجتمعة.",
            "red_flags": red_flags,
            "advice": "تعامل معه كتصيد ولا تشارك أي بيانات."
        }

    return None


def generate_safe_file_answer(filename):
    return {
        "answer_type": "analysis",
        "classification": "Safe",
        "confidence": 0.82,
        "reason": "تم تحليل محتوى الملف ولم تظهر مؤشرات تصيد واضحة مثل طلب كلمة مرور، OTP، بيانات بنكية، رابط تسجيل دخول مشبوه، تهديد، أو استعجال.",
        "red_flags": [],
        "advice": "يبدو الملف آمنًا من ناحية مؤشرات التصيد الظاهرة."
    }


def generate_unreadable_file_answer(filename, content_type, note):
    return {
        "answer_type": "analysis",
        "classification": "Needs Review",
        "confidence": 0.30,
        "reason": f"لم أستطع استخراج أو تحليل محتوى واضح من الملف {filename}. هذه حالة تقنية وليست حكمًا بأنه مشبوه.",
        "red_flags": [],
        "advice": "جرّب رفع نسخة أوضح من الملف أو صورة واضحة.",
        "technical_note": note,
        "content_type": content_type
    }


def _file_risk_texts(user_message):
    lang = detect_response_language(user_message)

    if lang == "English":
        return {
            "dangerous_extension": "Dangerous or executable file extension",
            "double_extension": "Double extension may hide the real file type",
            "macro_extension": "Macro-enabled Office document",
            "archive": "Compressed archive may hide malicious files",
            "pdf_javascript": "PDF contains JavaScript or automatic actions",
            "pdf_embedded": "PDF contains embedded files or launch actions",
            "office_macro": "Office file contains macro code",
            "office_external_links": "Office file contains external links",
            "large_file": "Unusually large file size",
            "svg_script": "SVG image may contain scripts or external references",
            "reason_intro": "The file content may look normal, but the file type or internal structure has risk indicators.",
            "safe_reason": "No risky file-type indicators were found before content analysis.",
            "advice": "Do not open risky files directly. Scan the file with antivirus or a sandbox first, and only open it from a trusted source.",
            "safe_advice": "The file structure looks normal, but still open files only from trusted sources.",
        }

    return {
        "dangerous_extension": "امتداد ملف تنفيذي أو خطير",
        "double_extension": "امتداد مزدوج قد يخفي نوع الملف الحقيقي",
        "macro_extension": "ملف Office يدعم الماكروز",
        "archive": "ملف مضغوط قد يخفي ملفات ضارة",
        "pdf_javascript": "ملف PDF يحتوي JavaScript أو إجراءات تلقائية",
        "pdf_embedded": "ملف PDF يحتوي ملفات مرفقة أو أوامر تشغيل",
        "office_macro": "ملف Office يحتوي كود Macro",
        "office_external_links": "ملف Office يحتوي روابط خارجية",
        "large_file": "حجم الملف غير معتاد وكبير",
        "svg_script": "صورة SVG قد تحتوي سكربتات أو مراجع خارجية",
        "reason_intro": "قد يبدو محتوى الملف طبيعيًا، لكن نوع الملف أو تركيبه الداخلي يحتوي مؤشرات خطورة.",
        "safe_reason": "لم تظهر مؤشرات خطورة من نوع الملف أو تركيبه قبل تحليل المحتوى.",
        "advice": "لا تفتح الملفات الخطيرة مباشرة. افحص الملف بمضاد فيروسات أو Sandbox أولًا، وافتحه فقط إذا كان من مصدر موثوق.",
        "safe_advice": "تركيبة الملف تبدو طبيعية، لكن افتح الملفات فقط من مصادر موثوقة.",
    }


def static_file_safety_check(filename, content_type, data, user_message=""):
    """
    Static file/image risk layer.
    It does NOT execute/open the uploaded file. It only checks filename, extension,
    MIME type, magic bytes, and safe structural indicators.
    """
    texts = _file_risk_texts(user_message)
    name = (filename or "unknown_file").strip()
    lower_name = name.lower()
    suffixes = [s.lower().lstrip(".") for s in Path(lower_name).suffixes]
    last_ext = suffixes[-1] if suffixes else ""

    red_flags = []
    severity = 0

    dangerous_extensions = {
        "exe", "dmg", "app", "pkg", "msi", "scr", "com", "bat", "cmd",
        "ps1", "vbs", "js", "jse", "jar", "apk", "sh", "run", "dll"
    }
    macro_extensions = {"docm", "xlsm", "pptm", "xlam", "dotm"}
    archive_extensions = {"zip", "rar", "7z", "tar", "gz", "bz2", "xz"}
    common_decoy_extensions = {
        "pdf", "doc", "docx", "xls", "xlsx", "ppt", "pptx", "txt",
        "jpg", "jpeg", "png", "webp", "csv", "html"
    }

    if last_ext in dangerous_extensions:
        red_flags.append(texts["dangerous_extension"])
        severity += 3

    if len(suffixes) >= 2 and suffixes[-2] in common_decoy_extensions and suffixes[-1] in dangerous_extensions:
        red_flags.append(texts["double_extension"])
        severity += 3

    if last_ext in macro_extensions:
        red_flags.append(texts["macro_extension"])
        severity += 2

    if last_ext in archive_extensions:
        red_flags.append(texts["archive"])
        severity += 1

    if len(data) > 50 * 1024 * 1024:
        red_flags.append(texts["large_file"])
        severity += 1

    if last_ext == "pdf" or content_type == "application/pdf" or data[:4] == b"%PDF":
        sample = data[:2_000_000]
        if any(token in sample for token in [b"/JavaScript", b"/JS", b"/OpenAction", b"/AA"]):
            red_flags.append(texts["pdf_javascript"])
            severity += 2
        if any(token in sample for token in [b"/EmbeddedFile", b"/Launch", b"/RichMedia"]):
            red_flags.append(texts["pdf_embedded"])
            severity += 2

    if last_ext in {"docx", "docm", "xlsx", "xlsm", "pptx", "pptm"} or data[:2] == b"PK":
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as z:
                names = [n.lower() for n in z.namelist()]
                if any(n.endswith("vbaproject.bin") for n in names):
                    red_flags.append(texts["office_macro"])
                    severity += 3
                rel_files = [n for n in z.namelist() if n.lower().endswith(".rels")]
                for rel in rel_files[:20]:
                    rel_text = z.read(rel)[:200000].decode("utf-8", errors="ignore").lower()
                    if "targetmode=\"external\"" in rel_text or "http://" in rel_text or "https://" in rel_text:
                        red_flags.append(texts["office_external_links"])
                        severity += 1
                        break
        except Exception:
            pass

    if last_ext == "svg" or content_type in {"image/svg+xml", "text/xml"}:
        sample_text = data[:500000].decode("utf-8", errors="ignore").lower()
        if "<script" in sample_text or "javascript:" in sample_text or "http://" in sample_text or "https://" in sample_text:
            red_flags.append(texts["svg_script"])
            severity += 2

    red_flags = list(dict.fromkeys(red_flags))

    if severity >= 3:
        classification = "Suspicious"
        confidence = 0.78
    elif severity >= 1:
        classification = "Suspicious"
        confidence = 0.62
    else:
        classification = "Safe"
        confidence = 0.70

    return {
        "answer_type": "analysis",
        "classification": classification,
        "confidence": confidence,
        "reason": texts["reason_intro"] if red_flags else texts["safe_reason"],
        "red_flags": red_flags,
        "advice": texts["advice"] if red_flags else texts["safe_advice"],
        "file_static_severity": severity,
        "file_static_checked": True
    }


def merge_file_risk_with_content_answer(content_answer, static_answer, user_message=""):
    """Merge static file risk with LLM/content analysis without hiding either signal."""
    if not isinstance(content_answer, dict):
        content_answer = {}
    if not isinstance(static_answer, dict):
        return content_answer

    static_flags = static_answer.get("red_flags", []) or []
    static_severity = int(static_answer.get("file_static_severity", 0) or 0)
    if static_severity <= 0:
        content_answer["file_static_risk"] = static_answer
        return content_answer

    current_class = content_answer.get("classification", "Safe")
    current_conf = float(content_answer.get("confidence", 0.5) or 0.5)

    if current_class in ["Safe", "Needs Review"]:
        content_answer["classification"] = "Suspicious"
        content_answer["confidence"] = max(current_conf, float(static_answer.get("confidence", 0.62)))
    elif current_class == "Suspicious":
        content_answer["confidence"] = max(current_conf, float(static_answer.get("confidence", 0.62)))

    existing_flags = content_answer.get("red_flags", []) or []
    content_answer["red_flags"] = list(dict.fromkeys(existing_flags + static_flags))

    lang = detect_response_language(user_message)
    old_reason = content_answer.get("reason", "")
    old_advice = content_answer.get("advice", "")

    if lang == "English":
        content_answer["reason"] = (old_reason + "\n\nFile risk note: " + static_answer.get("reason", "")).strip()
        content_answer["advice"] = (old_advice + "\n\nFile safety advice: " + static_answer.get("advice", "")).strip()
    else:
        content_answer["reason"] = (old_reason + "\n\nملاحظة حول خطورة الملف: " + static_answer.get("reason", "")).strip()
        content_answer["advice"] = (old_advice + "\n\nنصيحة أمان للملف: " + static_answer.get("advice", "")).strip()

    content_answer["file_static_risk"] = static_answer
    return content_answer


@traceable(
    name="Educational Cybersecurity Answer",
    run_type="llm",
    tags=["educational", "cybersecurity", "llm"]
)
def generate_educational_answer(req_message, history_text):
    system_prompt = """
You are a cybersecurity teacher.
Return ONLY valid JSON.

STRICT LANGUAGE RULE:
- Detect the language from the current user message.
- If the user message is English, reason, red_flags, advice, explanations, examples, tips, and summaries MUST be English.
- If the user message is Arabic, reason, red_flags, advice, explanations, examples, tips, and summaries MUST be Arabic.
- Do NOT mix languages in human-readable values.
- JSON keys must stay in English.

JSON schema:
{
  "answer_type": "educational",
  "title": "title in the user's language",
  "explanation": "detailed explanation in the user's language",
  "examples": ["example 1 in the user's language", "example 2 in the user's language", "example 3 in the user's language"],
  "practical_tips": ["tip 1 in the user's language", "tip 2 in the user's language", "tip 3 in the user's language"],
  "short_summary": "short summary in the user's language"
}
"""

    user_prompt = f"""
Previous conversation:
{history_text}

User question:
{req_message}

{language_instruction(req_message)}

Return only valid JSON.
"""

    raw = call_ab_chat_model(system_prompt, user_prompt, allow_phi=True)
    parsed = safe_json_parse(raw)

    # Ensure educational answers always have the full structure expected by Streamlit.
    if not isinstance(parsed, dict) or parsed.get("answer_type") != "educational":
        parsed = {
            "answer_type": "educational",
            "title": "Educational explanation" if detect_response_language(req_message) == "English" else "شرح تعليمي",
            "explanation": raw,
            "examples": [],
            "practical_tips": [],
            "short_summary": raw[:300]
        }

    parsed.setdefault("answer_type", "educational")
    parsed.setdefault("title", "Educational explanation" if detect_response_language(req_message) == "English" else "شرح تعليمي")
    parsed.setdefault("explanation", raw)
    parsed.setdefault("examples", [])
    parsed.setdefault("practical_tips", [])
    parsed.setdefault("short_summary", parsed.get("explanation", "")[:300])

    if not parsed.get("explanation"):
        parsed["explanation"] = raw

    return json.dumps(parsed, ensure_ascii=False)


@traceable(
    name="Uploaded Text File Analysis",
    run_type="llm",
    tags=["file-analysis", "text-file", "phishing-detection"]
)
def generate_text_file_analysis(message, filename, content_type, extracted_text, method):
    system_prompt = """
You are a careful cybersecurity phishing detector.
Analyze ONLY the provided file text.
Return ONLY valid JSON.

STRICT LANGUAGE RULE:
- Detect the language from the current user message.
- If the user message is English, reason, red_flags, advice, explanations, examples, tips, and summaries MUST be English.
- If the user message is Arabic, reason, red_flags, advice, explanations, examples, tips, and summaries MUST be Arabic.
- Do NOT mix languages in human-readable values.
- JSON keys must stay in English.

JSON schema:
{
  "answer_type": "analysis",
  "classification": "Safe" | "Suspicious" | "Phishing",
  "confidence": number between 0 and 1,
  "reason": "explanation in the user's language",
  "red_flags": ["red flag 1 in the user's language", "red flag 2 in the user's language"],
  "advice": "practical advice in the user's language"
}

STRICT RULES:
- Do NOT classify every uploaded file as suspicious.
- If the file is a normal CV, cover letter, homework, lecture, report, or assignment with no risky request, classify Safe.
- Phishing only if there is clear evidence: password, OTP, bank/card info, fake login, suspicious link, urgent account suspension, impersonation, or malware-like instruction.
- Suspicious only if there are real weak indicators.
"""

    user_prompt = f"""
User message:
{message}

{language_instruction(message)}

File name:
{filename}

File type:
{content_type}

Extraction method:
{method}

Extracted text:
{extracted_text}

Return only valid JSON.
"""

    return call_ab_chat_model(system_prompt, user_prompt, allow_phi=True)


@traceable(
    name="Uploaded Image or PDF Vision Analysis",
    run_type="llm",
    tags=["vision-analysis", "file-analysis", "image-pdf"]
)
def generate_vision_analysis(message, filename, content_type, image_urls):
    system_prompt = """
You are a careful cybersecurity phishing detector with vision ability.
Analyze ONLY the provided image/PDF/DOCX pages.
Return ONLY valid JSON.

STRICT LANGUAGE RULE:
- Detect the language from the current user message.
- If the user message is English, reason, red_flags, advice, explanations, examples, tips, and summaries MUST be English.
- If the user message is Arabic, reason, red_flags, advice, explanations, examples, tips, and summaries MUST be Arabic.
- Do NOT mix languages in human-readable values.
- JSON keys must stay in English.

JSON schema:
{
  "answer_type": "analysis",
  "classification": "Safe" | "Suspicious" | "Phishing",
  "confidence": number between 0 and 1,
  "reason": "explanation in the user's language",
  "red_flags": ["red flag 1 in the user's language", "red flag 2 in the user's language"],
  "advice": "practical advice in the user's language"
}

STRICT RULES:
- Do NOT classify every uploaded file as suspicious.
- If it is a normal CV, cover letter, homework, assignment, report, lecture slide, invoice, or certificate with no risky request, classify Safe.
- Phishing only if there is clear evidence: password request, OTP request, bank/card request, fake login, suspicious link, urgent account suspension, impersonation, or malware-like instruction.
- Suspicious only if there are real weak indicators.
"""

    content = [{
        "type": "text",
        "text": f"""
User message:
{message}

{language_instruction(message)}

File name:
{filename}

File type:
{content_type}

Analyze visible content only.
Return only valid JSON.
"""
    }]

    for url in image_urls:
        content.append({
            "type": "image_url",
            "image_url": {"url": url}
        })

    LAST_USED_MODEL.set(VISION_MODEL)
    LAST_AB_VARIANT.set("A-VISION")

    response = client.chat.completions.create(
        model=VISION_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": content}
        ]
    )

    return response.choices[0].message.content


@traceable(
    name="RAG Answer Generation",
    run_type="llm",
    tags=["rag", "llm-generation", "phishing-analysis"]
)
def generate_answer_with_rag(req_message, history_text, follow_up, retrieval_query, context):
    system_prompt = """
You are a cybersecurity assistant specialized in scam and phishing detection.
Return ONLY valid JSON.

STRICT LANGUAGE RULE:
- Detect the language from the current user message.
- If the user message is English, reason, red_flags, advice, explanations, examples, tips, and summaries MUST be English.
- If the user message is Arabic, reason, red_flags, advice, explanations, examples, tips, and summaries MUST be Arabic.
- Do NOT mix languages in human-readable values.
- JSON keys must stay in English.

JSON schema:
{
  "answer_type": "analysis",
  "classification": "Safe" | "Suspicious" | "Phishing",
  "confidence": number between 0 and 1,
  "reason": "explanation in the user's language",
  "red_flags": ["red flag 1 in the user's language", "red flag 2 in the user's language"],
  "advice": "practical advice in the user's language"
}
"""

    user_prompt = f"""
Previous conversation:
{history_text}

Current user message:
{req_message}

{language_instruction(req_message)}

Is follow-up:
{follow_up}

Message used for retrieval:
{retrieval_query}

Retrieved context:
{context}

Return only valid JSON.
"""

    return call_ab_chat_model(system_prompt, user_prompt, allow_phi=True)


@traceable(
    name="LLM Fallback Answer Generation",
    run_type="llm",
    tags=["llm-fallback", "phishing-analysis"]
)
def generate_answer_with_llm(req_message, history_text, follow_up, retrieval_query):
    system_prompt = """
You are a cybersecurity assistant specialized in scam and phishing detection.
Return ONLY valid JSON.

STRICT LANGUAGE RULE:
- Detect the language from the current user message.
- If the user message is English, reason, red_flags, advice, explanations, examples, tips, and summaries MUST be English.
- If the user message is Arabic, reason, red_flags, advice, explanations, examples, tips, and summaries MUST be Arabic.
- Do NOT mix languages in human-readable values.
- JSON keys must stay in English.

JSON schema:
{
  "answer_type": "analysis",
  "classification": "Safe" | "Suspicious" | "Phishing",
  "confidence": number between 0 and 1,
  "reason": "explanation in the user's language",
  "red_flags": ["red flag 1 in the user's language", "red flag 2 in the user's language"],
  "advice": "practical advice in the user's language"
}
"""

    user_prompt = f"""
Previous conversation:
{history_text}

Current user message:
{req_message}

{language_instruction(req_message)}

Is follow-up:
{follow_up}

Message/file being analyzed:
{retrieval_query}

Return only valid JSON.
"""

    return call_ab_chat_model(system_prompt, user_prompt, allow_phi=True)


def adjust_confidence(answer, retrieved, mode):
    classification = answer.get("classification", "Suspicious")
    confidence = float(answer.get("confidence", 0.5))

    top_score = retrieved[0]["score"] if retrieved else 0
    avg_score = sum(r["score"] for r in retrieved) / len(retrieved) if retrieved else 0

    if classification == "Phishing":
        confidence = min(max(confidence, 0.75), 0.95)
    elif classification == "Suspicious":
        confidence = min(max(confidence, 0.40), 0.75)
    elif classification == "Safe":
        confidence = min(max(confidence, 0.65), 0.90)
    elif classification == "Needs Review":
        confidence = min(max(confidence, 0.20), 0.45)

    answer["confidence"] = round(confidence, 2)
    answer["retrieval_top_score"] = round(top_score, 4)
    answer["retrieval_avg_score"] = round(avg_score, 4)
    answer["answer_mode"] = mode

    return answer


@app.get("/")
def home():
    return {
        "status": "running",
        "version": "LANGSMITH_MONITORING_ENABLED",
        "system_type": "hybrid-rag-llm-file-image-pdf-docx-vision",
        "chat_model": CHAT_MODEL,
        "vision_model": VISION_MODEL,
        "embedding_model": EMBEDDING_MODEL,
        "knowledge_items": len(knowledge_items),
        "rag_threshold": RAG_THRESHOLD,
        "langsmith_project": os.getenv("LANGSMITH_PROJECT", "LLMProject")
    }


@app.post("/chat")
@traceable(
    name="Scam Phishing RAG Chatbot",
    run_type="chain",
    tags=["scam-detection", "phishing", "rag-chatbot"]
)
async def chat(
    message: str = Form(""),
    chat_history: str = Form("[]"),
    ab_model: str = Form("auto"),
    files: Optional[List[UploadFile]] = File(None)
):
    start_time = time.time()
    request_id = str(uuid.uuid4())[:8]
    selected_model, ab_variant = choose_ab_model(ab_model)

    logger.info(
        f"request_id={request_id} | event=chat_request_started | "
        f"message_length={len(message)} | has_files={files is not None}"
    )

    try:
        parsed_history = json.loads(chat_history)
    except Exception as e:
        logger.warning(f"request_id={request_id} | event=chat_history_parse_failed | error={str(e)}")
        parsed_history = []

    history_text = build_history_text(parsed_history)

    uploaded_files_info = []
    file_items = []

    if files:
        for uploaded in files:
            data = await uploaded.read()
            filename = uploaded.filename or "unknown_file"
            content_type = uploaded.content_type or "application/octet-stream"

            logger.info(
                f"request_id={request_id} | event=file_uploaded | "
                f"filename={filename} | content_type={content_type} | size_bytes={len(data)}"
            )

            uploaded_files_info.append({
                "filename": filename,
                "content_type": content_type,
                "size_bytes": len(data)
            })

            static_risk = static_file_safety_check(filename, content_type, data, message)

            logger.info(
                f"request_id={request_id} | event=static_file_safety_checked | "
                f"filename={filename} | classification={static_risk.get('classification')} | "
                f"severity={static_risk.get('file_static_severity')} | "
                f"flags_count={len(static_risk.get('red_flags', []))}"
            )

            save_static_file_risk_event(
                filename=filename,
                content_type=content_type,
                static_risk=static_risk,
                request_id=request_id,
                user_message=message
            )

            if is_image_file(filename, content_type):
                encoded = base64.b64encode(data).decode("utf-8")
                image_url = f"data:{content_type};base64,{encoded}"

                file_items.append({
                    "filename": filename,
                    "content_type": content_type,
                    "text": "",
                    "image_urls": [image_url],
                    "method": "image_upload",
                    "note": "",
                    "static_risk": static_risk
                })

            else:
                extraction = extract_text_from_upload(filename, content_type, data)

                logger.info(
                    f"request_id={request_id} | event=file_extracted | "
                    f"filename={filename} | method={extraction.get('method')} | "
                    f"readable={extraction.get('readable')} | "
                    f"text_length={len(extraction.get('text', ''))} | "
                    f"image_pages={len(extraction.get('image_urls', []))}"
                )

                file_items.append({
                    "filename": filename,
                    "content_type": content_type,
                    "text": extraction.get("text", ""),
                    "image_urls": extraction.get("image_urls", []),
                    "method": extraction.get("method", "none"),
                    "note": extraction.get("note", ""),
                    "static_risk": static_risk
                })

    if uploaded_files_info:
        file_results = []

        for item in file_items:
            filename = item["filename"]
            content_type = item["content_type"]
            extracted_text = item.get("text", "").strip()
            image_urls = item.get("image_urls", [])
            method = item.get("method", "none")
            note = item.get("note", "")
            static_risk = item.get("static_risk", {})

            if extracted_text:
                rule_answer = detect_phishing_rules(extracted_text)

                if rule_answer:
                    parsed_answer = rule_answer
                    raw_answer = json.dumps(rule_answer, ensure_ascii=False)
                else:
                    raw_answer = generate_text_file_analysis(
                        message=message,
                        filename=filename,
                        content_type=content_type,
                        extracted_text=extracted_text,
                        method=method
                    )
                    parsed_answer = safe_json_parse(raw_answer)

                    if parsed_answer.get("classification") == "Suspicious" and not parsed_answer.get("red_flags"):
                        parsed_answer = generate_safe_file_answer(filename)
                        raw_answer = json.dumps(parsed_answer, ensure_ascii=False)

                parsed_answer = merge_file_risk_with_content_answer(parsed_answer, static_risk, message)
                raw_answer = json.dumps(parsed_answer, ensure_ascii=False)
                parsed_answer = adjust_confidence(parsed_answer, [], "FILE_TEXT_ANALYSIS")
                mode = "FILE_TEXT_ANALYSIS"
                preview = extracted_text[:700]

            elif image_urls:
                raw_answer = generate_vision_analysis(
                    message=message,
                    filename=filename,
                    content_type=content_type,
                    image_urls=image_urls
                )
                parsed_answer = safe_json_parse(raw_answer)
                parsed_answer = merge_file_risk_with_content_answer(parsed_answer, static_risk, message)
                raw_answer = json.dumps(parsed_answer, ensure_ascii=False)
                parsed_answer = adjust_confidence(parsed_answer, [], "VISION_FILE_ANALYSIS")
                mode = "VISION_FILE_ANALYSIS"
                preview = "[Analyzed visually]"

            else:
                parsed_answer = generate_unreadable_file_answer(
                    filename=filename,
                    content_type=content_type,
                    note=note
                )
                parsed_answer = merge_file_risk_with_content_answer(parsed_answer, static_risk, message)
                parsed_answer = adjust_confidence(parsed_answer, [], "UNREADABLE_FILE")
                raw_answer = json.dumps(parsed_answer, ensure_ascii=False)
                mode = "UNREADABLE_FILE"
                preview = ""

            logger.info(
                f"request_id={request_id} | event=file_analysis_result | "
                f"filename={filename} | mode={mode} | "
                f"classification={parsed_answer.get('classification')} | "
                f"confidence={parsed_answer.get('confidence')}"
            )

            file_results.append({
                "filename": filename,
                "content_type": content_type,
                "mode": mode,
                "answer": parsed_answer,
                "raw_answer": raw_answer,
                "extracted_preview": preview,
                "extraction_method": method,
                "extraction_note": note,
                "file_static_risk": static_risk
            })

            # Save each analyzed file/image into chat history so the dashboard counters,
            # static risk section, full logs, and high-risk events update after uploads.
            try:
                saved_message = message.strip() if message and message.strip() else "Uploaded file/image for security analysis"
                saved_message = f"{saved_message} | File: {filename}"
                save_chat_history(
                    user_message=saved_message,
                    bot_answer=json.dumps(parsed_answer, ensure_ascii=False),
                    mode=mode
                )
            except Exception as e:
                logger.error(f"event=file_chat_history_save_failed | error={str(e)}")

        latency = round(time.time() - start_time, 2)

        logger.info(
            f"request_id={request_id} | event=file_analysis_completed | "
            f"latency={latency}s | files_count={len(file_results)}"
        )

        # Save A/B/file analysis records for the admin dashboard only.
        for item in file_results:
            save_ab_test_result(
                user_message=message if message.strip() else "Uploaded file/image for security analysis",
                answer=item.get("answer", {}),
                mode=item.get("mode", "FILE_ANALYSIS"),
                model_used=item.get("answer", {}).get("model_used", LAST_USED_MODEL.get()),
                ab_variant=item.get("answer", {}).get("ab_variant", LAST_AB_VARIANT.get()),
                latency=latency,
                top_score=0,
                intent="file_or_image_analysis",
                request_id=request_id,
                filename=item.get("filename", "")
            )

        return {
            "mode": "MULTI_FILE_IMAGE_ANALYSIS",
            "system_type": "hybrid-rag-llm-file-image-pdf-docx-vision",
            "intent": "file_or_image_analysis",
            "latency": latency,
            "model_used": LAST_USED_MODEL.get(),
            "ab_variant": LAST_AB_VARIANT.get(),
            "ab_test_enabled": True,
            "file_results": file_results,
            "uploaded_files": uploaded_files_info,
            "retrieved_sources": []
        }

    intent = detect_user_intent(message)

    logger.info(f"request_id={request_id} | event=intent_detected | intent={intent}")

    if intent == "educational":
        raw_answer = generate_educational_answer(message, history_text)
        parsed_answer = safe_json_parse(raw_answer)
        parsed_answer = add_ab_metadata(parsed_answer)
        latency = round(time.time() - start_time, 2)

        logger.info(
            f"request_id={request_id} | event=educational_answer_completed | "
            f"latency={latency}s | intent={intent}"
        )

        save_ab_test_result(
            user_message=message,
            answer=parsed_answer,
            mode="LLM_EDUCATIONAL",
            model_used=LAST_USED_MODEL.get(),
            ab_variant=LAST_AB_VARIANT.get(),
            latency=latency,
            top_score=0,
            intent=intent,
            request_id=request_id
        )

        return {
            "mode": "LLM_EDUCATIONAL",
            "system_type": "hybrid-rag-llm-file-image-pdf-docx-vision",
            "intent": intent,
            "latency": latency,
            "model_used": LAST_USED_MODEL.get(),
            "ab_variant": LAST_AB_VARIANT.get(),
            "ab_test_enabled": True,
            "answer": parsed_answer,
            "raw_answer": raw_answer,
            "retrieved_sources": []
        }

    follow_up = is_follow_up(message)
    last_user_message = get_last_user_message(parsed_history)

    if follow_up and last_user_message:
        retrieval_query = last_user_message + "\n" + message
    else:
        retrieval_query = message

    retrieved = retrieve_by_embedding(retrieval_query, top_k=3)
    context = build_context(retrieved)
    top_score = retrieved[0]["score"] if retrieved else 0

    logger.info(
        f"request_id={request_id} | event=retrieval_completed | "
        f"top_score={round(top_score, 4)} | threshold={RAG_THRESHOLD} | follow_up={follow_up}"
    )

    if top_score >= RAG_THRESHOLD:
        mode = "RAG_ANALYSIS"
        raw_answer = generate_answer_with_rag(
            req_message=message,
            history_text=history_text,
            follow_up=follow_up,
            retrieval_query=retrieval_query,
            context=context
        )
    else:
        mode = "LLM_ANALYSIS"
        raw_answer = generate_answer_with_llm(
            req_message=message,
            history_text=history_text,
            follow_up=follow_up,
            retrieval_query=retrieval_query
        )

    parsed_answer = safe_json_parse(raw_answer)
    parsed_answer = adjust_confidence(parsed_answer, retrieved, mode)

    latency = round(time.time() - start_time, 2)

    logger.info(
        f"request_id={request_id} | event=chat_answer_completed | "
        f"mode={mode} | intent={intent} | "
        f"classification={parsed_answer.get('classification')} | "
        f"confidence={parsed_answer.get('confidence')} | "
        f"top_score={round(top_score, 4)} | latency={latency}s"
    )

    # =========================
    # SAVE CHAT FOR EVALUATION
    # =========================
    try:
        save_chat_history(
            user_message=message,
            bot_answer=json.dumps(parsed_answer, ensure_ascii=False),
            mode=mode
        )
    except Exception as e:
        logger.error(f"event=chat_history_save_failed | error={str(e)}")

    save_ab_test_result(
        user_message=message,
        answer=parsed_answer,
        mode=mode,
        model_used=LAST_USED_MODEL.get(),
        ab_variant=LAST_AB_VARIANT.get(),
        latency=latency,
        top_score=top_score,
        intent=intent,
        request_id=request_id
    )

    return {
        "mode": mode,
        "system_type": "hybrid-rag-llm-file-image-pdf-docx-vision",
        "intent": intent,
        "is_follow_up": follow_up,
        "latency": latency,
        "model_used": LAST_USED_MODEL.get(),
        "ab_variant": LAST_AB_VARIANT.get(),
        "ab_test_enabled": True,
        "rag_threshold": RAG_THRESHOLD,
        "answer": parsed_answer,
        "raw_answer": raw_answer,
        "retrieval_query": retrieval_query,
        "retrieved_sources": [
            {
                "score": round(result["score"], 4),
                "risk": result["item"].get("risk"),
                "text": result["item"].get("text", "")[:300],
                "advice": result["item"].get("advice")
            }
            for result in retrieved
        ]
    }